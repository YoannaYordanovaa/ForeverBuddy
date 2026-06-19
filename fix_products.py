import json
import os
import re
import sys

# Windows UTF-8 fix — принтиране на кирилица в терминала
if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if sys.stderr.encoding and sys.stderr.encoding.lower() != 'utf-8':
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

# ── КОНФИГУРАЦИЯ ────────────────────────────────────────────
INPUT_FOLDER = "input_files"
OUTPUT_FILE  = "output_md/products_optimized.md"

# Полета, които търсим (на БГ и EN) — нормализираме към стандартен ключ
FIELD_MAP = {
    # Име
    "име": "name", "ime": "name", "name": "name", "product": "name",
    "продукт": "name", "naziv": "name", "наименование": "name", "title": "name",
    # Цена
    "цена": "price", "tsena": "price", "price": "price", "cena": "price",
    "стойност": "price", "cost": "price",
    # Категория
    "категория": "category", "kategoriya": "category", "category": "category",
    "тип": "category", "type": "category", "вид": "category",
    # Описание
    "описание": "description", "opisanie": "description",
    "description": "description", "opis": "description",
    "info": "description", "информация": "description", "benefits": "description",
    "ползи": "description",
    # Съставки
    "състав": "ingredients", "sastav": "ingredients",
    "ingredients": "ingredients", "ingr": "ingredients",
    "composition": "ingredients", "съдържание": "ingredients",
    # Употреба
    "начин на употреба": "usage", "употреба": "usage", "usage": "usage",
    "upotreба": "usage", "how to use": "usage", "directions": "usage",
    "дозировка": "usage", "dosage": "usage",
    # Български ключове с интервали (от JSON)
    "начин на употреба": "usage",
    "описание и ползи": "description",
    "съставки": "ingredients",
    "ime": "name", "tsena": "price",
}

def normalize_key(raw_key: str) -> str:
    """Нормализира произволен ключ към стандартен."""
    k = " ".join(raw_key.strip().lower().split())  # trim + collapse whitespace
    return FIELD_MAP.get(k, k)

def normalize_price(raw: str) -> str:
    """Привежда цената към формат '00.00 лв. / 00.00 EUR'"""
    if not raw or str(raw).strip() in ("—", "-", "н/д", "", "None", "Цена при запитване"):
        return "Цена при запитване"
    raw = str(raw).strip()

    # Вече е форматирана на БГ
    if "лв" in raw.lower():
        return raw

    # Формат: "61.80 BGN / 31.60 EUR"
    import re as _re
    bgn_match = _re.search(r"([\d]+[.,][\d]+|[\d]+)\s*BGN", raw, _re.IGNORECASE)
    eur_match = _re.search(r"([\d]+[.,][\d]+|[\d]+)\s*EUR", raw, _re.IGNORECASE)

    if bgn_match:
        bgn = float(bgn_match.group(1).replace(",", "."))
        result = f"{bgn:.2f} лв."
        if eur_match:
            eur = float(eur_match.group(1).replace(",", "."))
            result += f" / {eur:.2f} EUR"
        return result

    if eur_match:
        eur = float(eur_match.group(1).replace(",", "."))
        return f"{eur:.2f} EUR"

    nums = _re.findall(r"[\d]+[.,]?[\d]*", raw)
    if nums:
        return f"{float(nums[0].replace(',', '.')):.2f} лв."

    return raw


def build_product_entry(p: dict, source: str) -> str:
    """Превръща речник с продукт в стандартен Markdown блок."""
    name        = p.get("name", "Неизвестен продукт").strip()
    bg_name     = (p.get("Българско_име") or p.get("българско_име") or
                   p.get("ime_bg") or p.get("Име_БГ") or "").strip()
    price       = normalize_price(str(p.get("price", "")))
    category    = p.get("category", "Общи продукти").strip()
    description = p.get("description", "Няма описание").replace("\n", " ").strip()
    ingredients = p.get("ingredients", "Няма данни за състава").strip()
    usage       = p.get("usage", "Виж етикета").strip()

    display_name = f"{name} / {bg_name}" if bg_name and bg_name != name else name

    return (
        f"=== ПРОДУКТ: {display_name} ===\n"
        f"ИМЕ: {name}\n"
        f"БЪЛГАРСКО ИМЕ: {bg_name if bg_name else name}\n"
        f"КАТЕГОРИЯ: {category}\n"
        f"ЦЕНА: {price}\n"
        f"СЪСТАВКИ: {ingredients}\n"
        f"НАЧИН НА УПОТРЕБА: {usage}\n"
        f"ОПИСАНИЕ И ПОЛЗИ: {description}\n"
        f"ИЗТОЧНИК: {source}\n"
        f"ТЕМА: продукт\n"
        f"{'─' * 50}\n\n"
    )


def parse_json(filepath: str) -> list[dict]:
    """Парсва JSON файл с продукти."""
    with open(filepath, "r", encoding="utf-8") as f:
        data = json.load(f)
    items = data if isinstance(data, list) else data.get("products", [data])
    result = []
    for item in items:
        normalized = {}
        for k, v in item.items():
            std_key = normalize_key(k)
            normalized[std_key] = v
        result.append(normalized)
    return result

def parse_excel(filepath: str) -> list[dict]:
    """Парсва Excel (.xlsx/.xls) файл."""
    try:
        import openpyxl
    except ImportError:
        print("  ⚠️  Инсталирай: pip install openpyxl")
        return []

    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    headers = [normalize_key(str(h)) if h else f"col_{i}" for i, h in enumerate(rows[0])]
    result = []
    for row in rows[1:]:
        if not any(row):
            continue
        product = {}
        for h, val in zip(headers, row):
            if val is not None:
                product[h] = str(val).strip()
        result.append(product)
    return result

def parse_word(filepath: str) -> list[dict]:
    """
    Парсва Word (.docx) файл.
    Поддържа два формата:
      А) Таблица с хедъри в първи ред
      Б) Параграфи с формат "Поле: Стойност"
    """
    try:
        from docx import Document
    except ImportError:
        print("  ⚠️  Инсталирай: pip install python-docx")
        return []

    doc = Document(filepath)
    result = []

    # Формат А: таблици
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [normalize_key(cell.text) for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            product = {}
            for h, cell in zip(headers, row.cells):
                txt = cell.text.strip()
                if txt:
                    product[h] = txt
            if product:
                result.append(product)

    # Формат Б: параграфи "Поле: Стойност" (само ако таблици не са намерени)
    if not result:
        current = {}
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                if current:
                    result.append(current)
                    current = {}
                continue
            if ":" in text:
                parts = text.split(":", 1)
                key = normalize_key(parts[0])
                val = parts[1].strip()
                current[key] = val
            elif current.get("name"):
                # Добавяме към описанието ако няма ключ
                current["description"] = current.get("description", "") + " " + text
        if current:
            result.append(current)

    return result

def parse_pdf(filepath: str) -> list[dict]:
    """
    Парсва PDF файл.
    Търси блокове от вида:
      Продукт: X
      Цена: Y
      ...
    или таблични структури.
    """
    try:
        import pdfplumber
    except ImportError:
        print("  ⚠️  Инсталирай: pip install pdfplumber")
        return []

    result = []
    current = {}

    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            # Опит за таблица
            tables = page.extract_tables()
            for table in tables:
                if not table or len(table) < 2:
                    continue
                headers = [normalize_key(str(h or "")) for h in table[0]]
                for row in table[1:]:
                    if not any(row):
                        continue
                    product = {}
                    for h, val in zip(headers, row):
                        if val:
                            product[h] = str(val).strip()
                    if product:
                        result.append(product)

            # Ако няма таблици — четем текст ред по ред
            if not tables:
                text = page.extract_text() or ""
                for line in text.split("\n"):
                    line = line.strip()
                    if not line:
                        if current:
                            result.append(current)
                            current = {}
                        continue
                    if ":" in line:
                        parts = line.split(":", 1)
                        key = normalize_key(parts[0])
                        val = parts[1].strip()
                        if key in FIELD_MAP.values():
                            current[key] = val

    if current:
        result.append(current)

    return result

# ── ГЛАВНА ФУНКЦИЯ ───────────────────────────────────────────

PARSERS = {
    ".json": parse_json,
    ".xlsx": parse_excel,
    ".xls":  parse_excel,
    ".docx": parse_word,
    ".doc":  parse_word,
    ".pdf":  parse_pdf,
}

def process_all_files():
    if not os.path.exists(INPUT_FOLDER):
        print(f"Грешка: Папката '{INPUT_FOLDER}' не съществува.")
        return

    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)

    files = [
        f for f in os.listdir(INPUT_FOLDER)
        if os.path.isfile(os.path.join(INPUT_FOLDER, f)) and not f.startswith(".")
    ]

    if not files:
        print("Папката е празна.")
        return

    all_products = []
    seen_names   = set()

    for filename in files:
        ext = os.path.splitext(filename)[1].lower()
        parser = PARSERS.get(ext)

        if not parser:
            print(f"  ⏭  Пропускам неподдържан формат: {filename}")
            continue

        filepath = os.path.join(INPUT_FOLDER, filename)
        print(f"Обработва се: {filename} ...", end=" ")

        try:
            products = parser(filepath)
            # Дедупликация по име
            added = 0
            for p in products:
                name = p.get("name", "").strip().lower()
                if name and name not in seen_names:
                    seen_names.add(name)
                    p["_source"] = filename
                    all_products.append(p)
                    added += 1
                elif name in seen_names:
                    print(f"\n  ↩  Дубликат пропуснат: {p.get('name')}", end="")
            print(f"✓ ({added} продукта)")
        except Exception as e:
            print(f"\n  ✗ Грешка: {e}")

    if not all_products:
        print("\nНяма намерени продукти.")
        return

    print(f"\nОбщо уникални продукти: {len(all_products)}")
    print(f"Записване в '{OUTPUT_FILE}'...")

    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        f.write("# ПРОДУКТОВ КАТАЛОГ — FOREVER LIVING PRODUCTS\n\n")
        f.write(f"Общо продукти: {len(all_products)}\n\n")
        f.write("=" * 60 + "\n\n")

        for p in all_products:
            source = p.pop("_source", "неизвестен")
            f.write(build_product_entry(p, source))

    print("✅ Готово! Стартирай create_db.py за да обновиш базата.")

if __name__ == "__main__":
    process_all_files()